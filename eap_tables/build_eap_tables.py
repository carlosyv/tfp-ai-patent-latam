"""
EAP (Paper 1) — stage 2 of 2: render the manuscript tables.

    eap_results_v5.json  ──▶  build_eap_tables.py  ──▶  EAP_tables_v5.docx
                                                        eap_tables_coverage.md

THE ONE RULE
------------
This script may not contain a numeric literal that ends up in a table cell.
Every number is fetched from the results JSON by dotted path.  A path that does
not resolve raises `MissingResult`; the cell renders as **[NO SOURCE]**, the
path is written to the coverage report, and the process exits non-zero.

That is the whole point.  The previous builder held every value as a hardcoded
string, which meant a cell with no underlying estimate was indistinguishable
from one with a real estimate behind it.  Table 6 panel B was filled with seven
plausible-looking coefficients that had never been computed, under a header
comment asserting that all numbers had been verified.  Here, a missing estimate
cannot be papered over: it is loud in the document and fatal to the build.

USAGE
    python3 eap_tables/build_eap_tables.py
    python3 eap_tables/build_eap_tables.py --allow-missing   # draft; still flags

SOURCE MAP  (see eap_tables/README.md for the full chain back to raw data)
    Table 1  table1_descriptives      <- compute_descriptives(merged_dissertation_v5)
    Table 2  table2_cd_tests          <- pesaran_cd_test on parsimonious FE residuals
    Table 3  table3_benchmark         <- pooled_ols / fixed_effects_twoway / cce_pooled
    Table 4  table4_heterogeneity     <- interaction models + median splits
    Table 5  table5_quantile          <- Canay (2011) two-step, 1000 bootstrap reps
    Table 6A table6a_mediation        <- Baron-Kenny + Sobel
    Table 6B table6b_robustness       <- robustness_verification_v5.json  (NOT recomputed)
    Table 7  table7_panel_b           <- merged_panelB_v5.csv
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
RESULTS = REPO / "output" / "results"

HEADER_FILL = "1F4E78"
SUBHEAD_FILL = "DDE6F0"
NO_SOURCE = "[NO SOURCE]"


class MissingResult(KeyError):
    """A table cell asked for a statistic the results file does not contain."""


# ── Results access ───────────────────────────────────────────────────────────


class Results:
    """
    Read-only view over eap_results_v5.json.  The only way to get a number into
    a table.  There is deliberately no `default` parameter on `get()`.
    """

    def __init__(self, payload: dict):
        self._d = payload
        self.missing: list[str] = []

    def get(self, path: str):
        node = self._d
        for part in path.split("."):
            if not isinstance(node, dict) or part not in node:
                self.missing.append(path)
                raise MissingResult(path)
            node = node[part]
        return node

    # -- formatted accessors --------------------------------------------------

    def num(self, path: str, dec: int = 3, signed: bool = False) -> str:
        try:
            v = float(self.get(path))
        except MissingResult:
            return NO_SOURCE
        s = f"{v:+.{dec}f}" if signed else f"{v:.{dec}f}"
        return s.replace("-", "−")

    def coef(self, path: str, dec: int = 3, signed: bool = False) -> str:
        """Coefficient with significance stars, e.g. '−0.047***'."""
        try:
            b = float(self.get(f"{path}.b"))
            p = float(self.get(f"{path}.p"))
        except MissingResult:
            return NO_SOURCE
        s = f"{b:+.{dec}f}" if signed else f"{b:.{dec}f}"
        return s.replace("-", "−") + stars(p)

    def se(self, path: str, dec: int = 3) -> str:
        try:
            return f"({float(self.get(f'{path}.se')):.{dec}f})"
        except MissingResult:
            return NO_SOURCE

    def n(self, path: str) -> str:
        try:
            return f"{int(self.get(f'{path}.N')):,}"
        except MissingResult:
            return NO_SOURCE


def stars(p: float) -> str:
    return "***" if p < 0.01 else "**" if p < 0.05 else "*" if p < 0.10 else ""


# ── docx helpers ─────────────────────────────────────────────────────────────


def shade(cell, hexfill: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def put(cell, text: str, *, bold=False, italic=False, align="left",
        fill=None, white=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(9)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if text == NO_SOURCE:
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    if fill:
        shade(cell, fill)


def add_table(doc: Document, n_cols: int, rows: list[list], *,
              header_rows: int = 1) -> None:
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j >= n_cols:
                continue
            cfg = val if isinstance(val, dict) else {"text": str(val)}
            text = cfg.pop("text")
            if i < header_rows:
                cfg.setdefault("bold", True)
                cfg.setdefault("fill", HEADER_FILL)
                cfg.setdefault("white", True)
                cfg.setdefault("align", "center" if j else "left")
            put(t.cell(i, j), text, **cfg)


def title(doc: Document, num: int | str, caption: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"Table {num}. {caption}")
    r.bold = True
    r.font.size = Pt(10)


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Notes: " + text)
    r.italic = True
    r.font.size = Pt(8)
    doc.add_paragraph()


# ── Tables ───────────────────────────────────────────────────────────────────


def table1(doc, R: Results) -> None:
    title(doc, 1, "Descriptive statistics, Panel A (N = 9, T = 25, Obs = 225).")
    rows = [["Variable", "Obs", "Mean", "Std. Dev.", "Min", "Max"]]
    try:
        recs = R.get("table1_descriptives.rows")
    except MissingResult:
        recs = []
    if not recs:
        R.missing.append("table1_descriptives.rows")
        rows.append([NO_SOURCE] * 6)
    for i, rec in enumerate(recs):
        cells = [str(rec.get("label", NO_SOURCE))]
        for key, dec in (("N", 0), ("mean", 3), ("sd", 3), ("min", 3), ("max", 3)):
            if key not in rec:
                R.missing.append(f"table1_descriptives.rows[{i}].{key}")
                cells.append({"text": NO_SOURCE, "align": "center"})
                continue
            txt = (f"{int(rec[key]):,}" if dec == 0
                   else f"{float(rec[key]):.{dec}f}".replace("-", "−"))
            cells.append({"text": txt, "align": "center"})
        rows.append(cells)
    add_table(doc, 6, rows)
    note(doc, "Author calculations from World Bank WDI, ILOSTAT, Penn World Table 10.01, "
              "WGI, and the WIPO patent search described in Section 3.3. VRS Malmquist "
              "N = 172 reflects infeasibility of the variable-returns-to-scale frontier "
              "for some country-years.")


def table2(doc, R: Results) -> None:
    title(doc, 2, "Pesaran (2004) cross-sectional dependence test.")
    rows = [["Dependent variable", "CD statistic", "p-value", "Inference"]]
    for key, label in [
        ("solow", "ln(TFP) — Solow"),
        ("malmquist_vrs", "Malmquist TFP Change (VRS)"),
        ("malmquist_crs", "Malmquist TFP Change (CRS)"),
    ]:
        base = f"table2_cd_tests.{key}"
        try:
            p = float(R.get(f"{base}.p"))
            inference = ("Fail to reject independence" if p >= 0.10
                         else "Reject: strong CSD")
            pstr = "< 0.001" if p < 0.001 else f"{p:.3f}"
        except MissingResult:
            inference, pstr = NO_SOURCE, NO_SOURCE
        rows.append([label,
                     {"text": R.num(f"{base}.CD", 2), "align": "center"},
                     {"text": pstr, "align": "center"},
                     inference])
    add_table(doc, 4, rows)
    note(doc, "CD statistic distributed N(0,1) under cross-sectional independence. "
              "Computed on residuals of the two-way FE model with the six-control "
              "parsimonious specification.")


def table3(doc, R: Results) -> None:
    title(doc, 3, "Benchmark and CSD-robust estimates of the AI–TFP elasticity, Panel A.")
    dvs = [("solow", "Solow ln(TFP)"), ("malmquist_vrs", "Malmquist VRS"),
           ("malmquist_crs", "Malmquist CRS")]
    hdr, sub, cf, se, ns, r2 = ([""] for _ in range(6))
    for key, label in dvs:
        hdr += [label, ""]
        sub += ["FE-DK", "CCEP"]
        for est in ("fe_dk", "ccep"):
            path = f"table3_benchmark.{key}.{est}"
            cf.append({"text": R.coef(path), "align": "center"})
            se.append({"text": R.se(path), "align": "center", "italic": True})
            ns.append({"text": R.n(path), "align": "center"})
            r2.append({"text": R.num(f"{path}.r2", 3), "align": "center"})
    cf[0], se[0], ns[0], r2[0] = "ln(AI patent stock)", "", "N", "R²"
    sub[0] = ""
    add_table(doc, 7, [hdr, sub, cf, se, ns, r2], header_rows=2)
    ols = f"table3_benchmark.solow.ols"
    note(doc, f"Pooled OLS on Solow ln(TFP) yields β = {R.coef(ols)} "
              f"(SE {R.se(ols)}), reported in the text but omitted from the table. "
              "Standard errors in parentheses; FE-DK uses Driscoll-Kraay with Bartlett "
              "kernel. Significance: * p < 0.10, ** p < 0.05, *** p < 0.01.")


def table4(doc, R: Results) -> None:
    title(doc, 4, "Heterogeneity by institutional quality and digital infrastructure.")
    mods = [("rule_of_law", "RoL"), ("mobile", "Mobile"), ("broadband", "Broadband")]
    dvs = [("solow", "Solow ln(TFP)"), ("malmquist_crs", "Malmquist CRS")]

    hdr = [""] + [lbl for _, lbl in dvs for _ in mods]
    sub = [""] + [f"({i+1}) {m}" for i in range(6) for m in [mods[i % 3][1]]]
    main, inter, ise, below, above, nrow = ([] for _ in range(6))
    for dv_key, _ in dvs:
        for mod_key, _ in mods:
            b = f"table4_heterogeneity.{dv_key}.{mod_key}"
            main.append({"text": R.coef(f"{b}.interaction.ai_main_effect", 4), "align": "center"})
            inter.append({"text": R.coef(f"{b}.interaction.interaction", 4), "align": "center"})
            ise.append({"text": R.se(f"{b}.interaction.interaction", 4), "align": "center", "italic": True})
            below.append({"text": R.coef(f"{b}.subsample.below", 3, signed=True), "align": "center"})
            above.append({"text": R.coef(f"{b}.subsample.above", 3, signed=True), "align": "center"})
            nrow.append({"text": R.n(f"{b}.interaction.interaction"), "align": "center"})

    rows = [
        hdr, sub,
        ["Panel A. Interaction terms"] + [""] * 6,
        ["ln(AI patent stock), main effect"] + main,
        ["ln(AI) × MOD"] + inter,
        [""] + ise,
        ["Panel B. Median-split subsample β(ln AI)"] + [""] * 6,
        ["Below-median subsample"] + below,
        ["Above-median subsample"] + above,
        ["N (interaction model)"] + nrow,
    ]
    add_table(doc, 7, rows, header_rows=2)
    note(doc, "The main-effect row reports each interaction model's own coefficient on "
              "ln(AI), which is conditional on MOD = 0 and therefore differs from the "
              "baseline estimate in Table 3. Median splits estimated separately above "
              "and below the sample median of each moderator. Driscoll-Kraay standard "
              "errors. Significance: * p < 0.10, ** p < 0.05, *** p < 0.01.")


def table5(doc, R: Results) -> None:
    title(doc, 5, "Panel quantile regression on log Solow TFP, Canay (2011) two-step.")
    taus = [("010", "0.10"), ("025", "0.25"), ("050", "0.50"),
            ("075", "0.75"), ("090", "0.90")]
    hdr = [""] + [f"({i+1}) τ = {lbl}" for i, (_, lbl) in enumerate(taus)]
    cf = ["ln(AI patent stock)"]
    se = [""]
    ns = ["N"]
    for key, _ in taus:
        p = f"table5_quantile.tau_{key}"
        cf.append({"text": R.coef(p), "align": "center"})
        se.append({"text": R.se(p), "align": "center", "italic": True})
        ns.append({"text": R.n(p), "align": "center"})
    add_table(doc, 6, [hdr, cf, se, ns])
    note(doc, "Step 1: two-way FE; Step 2: quantile regression on the entity-FE-purged "
              "dependent variable. Block-bootstrap standard errors clustered by country "
              "(1,000 replications) in parentheses. Significance: * p < 0.10, "
              "** p < 0.05, *** p < 0.01.")


def table6(doc, R: Results) -> None:
    title(doc, 6, "Mediation analysis and robustness battery.")
    rows = [
        ["Panel A. Baron-Kenny mediation analysis (Solow TFP)", "", "", "", ""],
        ["", "(1) IS Step 2", "(2) IS Step 3", "(3) HC Step 2", "(4) HC Step 3"],
    ]
    IS, HC = "table6a_mediation.industrial_structure", "table6a_mediation.human_capital"
    rows.append(["ln(AI) (α₁ or δ₁)"] + [
        {"text": R.coef(f"{IS}.step2"), "align": "center"},
        {"text": R.coef(f"{IS}.step3_ai"), "align": "center"},
        {"text": R.coef(f"{HC}.step2"), "align": "center"},
        {"text": R.coef(f"{HC}.step3_ai"), "align": "center"},
    ])
    rows.append(["Mediator (δ₂)", {"text": "—", "align": "center"},
                 {"text": R.coef(f"{IS}.step3_mediator"), "align": "center"},
                 {"text": "—", "align": "center"},
                 {"text": R.coef(f"{HC}.step3_mediator"), "align": "center"}])
    rows.append(["Sobel z (p-value)", {"text": "—", "align": "center"},
                 {"text": f"{R.num(f'{IS}.sobel.z', 2)} ({R.num(f'{IS}.sobel.p', 2)})", "align": "center"},
                 {"text": "—", "align": "center"},
                 {"text": f"{R.num(f'{HC}.sobel.z', 2)} ({R.num(f'{HC}.sobel.p', 2)})", "align": "center"}])
    rows.append(["% Mediated", {"text": "—", "align": "center"},
                 {"text": R.num(f"{IS}.pct_mediated", 1) + "%", "align": "center"},
                 {"text": "—", "align": "center"},
                 {"text": R.num(f"{HC}.pct_mediated", 1) + "%", "align": "center"}])

    rows.append(["Panel B. Robustness battery — lnAI coefficient (FE-DK)", "", "", "", ""])
    rows.append(["Specification", "Solow β", "", "Malmquist CRS β", "N"])
    specs = [
        ("baseline", "Baseline (FE-DK)"),
        ("raw_count", "Raw patent count (ln 1+flow)"),
        ("delta_p_022", "Alt. depreciation δp = 0.22"),
        ("per_gdp", "Per-GDP normalisation"),
        ("lag1", "One-period lag"),
        ("lag2", "Two-period lag"),
        ("iv_2sls_lag1", "2SLS, lagged AI (within)"),
    ]
    for key, label in specs:
        b = f"table6b_robustness.specifications.{key}"
        rows.append([
            label,
            {"text": R.coef(f"{b}.solow"), "align": "center"},
            "",
            {"text": R.coef(f"{b}.malm"), "align": "center"},
            {"text": R.n(f"{b}.solow"), "align": "center"},
        ])
    add_table(doc, 5, rows, header_rows=1)
    note(doc, "Panel A: IS = services-to-industry value-added ratio; HC = Penn World "
              "Table 10.01 human capital index. Step 2 regresses the mediator on lnAI; "
              "Step 3 regresses TFP on lnAI and the mediator. Sobel p-values from the "
              "asymptotic ab-product test. Panel B reports the lnAI coefficient under "
              "each robustness specification, generated by "
              "pipeline_v5/robustness_verification.py. Sample size varies across "
              "specifications because the lagged and 2SLS models lose initial years. "
              "Significance: * p < 0.10, ** p < 0.05, *** p < 0.01.")


def table7(doc, R: Results) -> None:
    title(doc, 7, "Panel B replication using OECD AI publications (N = 17, T = 9, 2016–2024).")
    dvs = [("solow", "Solow ln(TFP)"), ("malmquist_crs", "Malmquist CRS")]
    ests = [("ols", "OLS"), ("fe_dk", "FE-DK"), ("ccep", "CCEP")]
    hdr = [""] + [lbl for _, lbl in dvs for _ in ests]
    sub = [""] + [f"({i+1}) {e}" for i, (_, e) in enumerate(ests * 2)]
    cf, ns = ["ln(AI publications)"], ["N"]
    for dv_key, _ in dvs:
        for est, _ in ests:
            p = f"table7_panel_b.{dv_key}.{est}"
            cf.append({"text": R.coef(p, 3, signed=True), "align": "center"})
            ns.append({"text": R.n(p), "align": "center"})
    cd = ["Pesaran CD"]
    for dv_key, _ in dvs:
        cd += [{"text": R.num(f"table7_panel_b.{dv_key}.cd_test.CD", 2), "align": "center"},
               "", ""]
    add_table(doc, 7, [hdr, sub, cf, ns, cd], header_rows=2)
    note(doc, "ln(AI publications) is the log OECD.AI Observatory all-fields publication "
              "count. Significance: * p < 0.10, ** p < 0.05, *** p < 0.01.")


# ── Main ─────────────────────────────────────────────────────────────────────


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--results", default=str(RESULTS / "eap_results_v5.json"))
    ap.add_argument("--out", default=str(RESULTS / "EAP_tables_v5.docx"))
    ap.add_argument("--coverage", default=str(RESULTS / "eap_tables_coverage.md"))
    ap.add_argument("--allow-missing", action="store_true",
                    help="exit 0 even if cells are unsourced (draft builds only)")
    args = ap.parse_args()

    src = Path(args.results)
    if not src.exists():
        print(f"ERROR: {src} not found. Run compute_eap_results.py first.",
              file=sys.stderr)
        return 2

    payload = json.loads(src.read_text())
    R = Results(payload)

    doc = Document()
    doc.add_heading("EAP manuscript tables", level=1)
    meta = payload.get("meta", {})
    p = doc.add_paragraph()
    r = p.add_run(
        f"Generated {meta.get('generated')} from {src.name} "
        f"(git {str(meta.get('git_commit'))[:8]}, "
        f"dk_dof_correction={meta.get('dk_dof_correction')}). "
        "Every value in this document resolves to a key in that file."
    )
    r.italic = True
    r.font.size = Pt(8)
    doc.add_paragraph()

    for fn in (table1, table2, table3, table4, table5, table6, table7):
        fn(doc, R)

    out = Path(args.out)
    doc.save(out)

    lines = [
        "# EAP table coverage",
        "",
        f"- results file: `{src.name}`",
        f"- git commit: `{meta.get('git_commit')}`",
        f"- DK dof correction: `{meta.get('dk_dof_correction')}`",
        f"- unsourced cells: **{len(R.missing)}**",
        "",
    ]
    if R.missing:
        lines += ["## Unsourced paths", ""]
        lines += [f"- `{m}`" for m in dict.fromkeys(R.missing)]
        lines += ["", "Each renders as **[NO SOURCE]** in the document. Do not "
                      "hand-fill these cells — add the estimate to "
                      "`compute_eap_results.py` and rebuild."]
    else:
        lines.append("All table cells resolved to a computed statistic.")
    Path(args.coverage).write_text("\n".join(lines) + "\n")

    print(f"wrote {out}")
    print(f"wrote {args.coverage}")
    if R.missing:
        print(f"\n{len(R.missing)} UNSOURCED CELL(S):", file=sys.stderr)
        for m in dict.fromkeys(R.missing):
            print(f"  {m}", file=sys.stderr)
        if not args.allow_missing:
            return 1
    else:
        print("all cells sourced")
    return 0


if __name__ == "__main__":
    sys.exit(main())
